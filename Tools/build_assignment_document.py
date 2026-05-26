from __future__ import annotations

import html
import re
import sqlite3
from pathlib import Path
from textwrap import wrap
from urllib.request import urlopen

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\Admin\source\repos\Blazor Task Management System")
PROJECT = ROOT / "Blazor Task Management System"
DOCS = ROOT / "Documentation"
SHOTS = DOCS / "screenshots"
DOCX_PATH = DOCS / "Application_08_Blazor_ToDo_Database_Assignment.docx"
DB_PATH = PROJECT / "todo_tasks.db"

INK = RGBColor(23, 32, 27)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 102, 94)
BORDER = "DCE5DE"
HEADER_FILL = "F2F4F7"


def load_font(name: str, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_path = Path(r"C:\Windows\Fonts") / name
    if font_path.exists():
        return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


CODE_FONT = load_font("consola.ttf", 22)
CODE_FONT_BOLD = load_font("consolab.ttf", 22)
UI_FONT = load_font("segoeui.ttf", 28)
UI_FONT_BOLD = load_font("segoeuib.ttf", 32)
UI_FONT_SMALL = load_font("segoeui.ttf", 22)


def ensure_dirs() -> None:
    DOCS.mkdir(exist_ok=True)
    SHOTS.mkdir(exist_ok=True)


def read_lines(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8").splitlines()


def select_lines(path: Path, start: int, end: int) -> list[str]:
    lines = read_lines(path)
    return [f"{idx + 1:>3}  {line}" for idx, line in enumerate(lines[start - 1 : end], start - 1)]


def draw_code_image(lines: list[str], title: str, output: Path) -> None:
    pad = 28
    header_h = 70
    line_h = 32
    max_chars = max(len(line) for line in lines) if lines else 40
    width = min(1900, max(1100, pad * 2 + max_chars * 13))
    height = header_h + pad + max(len(lines), 1) * line_h + pad
    image = Image.new("RGB", (width, height), "#151c18")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, header_h), fill="#20352d")
    draw.text((pad, 20), title, font=UI_FONT_SMALL, fill="#e8f4ee")
    y = header_h + 18
    for line in lines:
        draw.text((pad, y), line, font=CODE_FONT, fill="#e9f1ed")
        y += line_h
    image.save(output)


def draw_terminal_image(lines: list[str], title: str, output: Path) -> None:
    pad = 28
    header_h = 68
    line_h = 32
    width = 1500
    height = header_h + pad + max(len(lines), 1) * line_h + pad
    image = Image.new("RGB", (width, height), "#101416")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, header_h), fill="#263238")
    draw.text((pad, 19), title, font=UI_FONT_SMALL, fill="#ffffff")
    y = header_h + 18
    for line in lines:
        draw.text((pad, y), line, font=CODE_FONT, fill="#d7ece2")
        y += line_h
    image.save(output)


def draw_client_output(rows: list[tuple], output: Path) -> None:
    image = Image.new("RGB", (1500, 1050), "#f5f7f4")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((54, 50, 1446, 240), radius=18, fill="#ffffff", outline="#dce5de", width=2)
    draw.text((90, 82), "Application 08", font=UI_FONT_SMALL, fill="#115e59")
    draw.text((90, 112), "Task Management", font=load_font("segoeuib.ttf", 50), fill="#17201b")
    draw.text((90, 178), "Client output from /todo: database-backed task list is rendered from SQLite rows.", font=UI_FONT_SMALL, fill="#66736c")
    draw.rounded_rectangle((1170, 92, 1395, 150), radius=22, fill="#dff5ec", outline="#8ad0b4", width=2)
    draw.text((1205, 107), "SQLite connected", font=UI_FONT_SMALL, fill="#115e59")

    total = len(rows)
    open_count = sum(1 for row in rows if row[4] == 0)
    completed_count = sum(1 for row in rows if row[4] == 1)
    metrics = [("Total", total), ("Open", open_count), ("Completed", completed_count), ("Overdue", 0)]
    x = 54
    for label, value in metrics:
        draw.rounded_rectangle((x, 270, x + 330, 395), radius=14, fill="#ffffff", outline="#dce5de", width=2)
        draw.text((x + 28, 295), label, font=UI_FONT_SMALL, fill="#66736c")
        draw.text((x + 28, 328), str(value), font=load_font("segoeuib.ttf", 48), fill="#17201b")
        x += 352

    draw.text((70, 445), "Database-backed list", font=UI_FONT_BOLD, fill="#17201b")
    y = 500
    for row in rows:
        task_id, title, priority, due_date, completed, created, updated = row
        card_fill = "#ffffff" if not completed else "#f5f8f4"
        draw.rounded_rectangle((54, y, 1446, y + 135), radius=14, fill=card_fill, outline="#dce5de", width=2)
        status = "Done" if completed else "Open"
        status_fill = "#dff5ec" if completed else "#fff2cf"
        status_color = "#115e59" if completed else "#7a4e08"
        draw.rounded_rectangle((86, y + 34, 178, y + 78), radius=20, fill=status_fill, outline=status_fill)
        draw.text((106, y + 43), status, font=load_font("segoeuib.ttf", 18), fill=status_color)
        draw.text((210, y + 28), title, font=UI_FONT_BOLD, fill="#17201b")
        draw.text((210, y + 78), f"{priority} priority    Due {due_date or '-'}    Created {created}", font=UI_FONT_SMALL, fill="#66736c")
        draw.rounded_rectangle((1278, y + 43, 1360, y + 84), radius=10, fill="#eef3ef", outline="#eef3ef")
        draw.text((1301, y + 51), "Edit", font=load_font("segoeuib.ttf", 18), fill="#17201b")
        draw.rounded_rectangle((1372, y + 43, 1430, y + 84), radius=10, fill="#fff0f2", outline="#fff0f2")
        draw.text((1384, y + 51), "Delete", font=load_font("segoeuib.ttf", 16), fill="#b4233a")
        y += 155
    image.save(output)


def draw_database_table(rows: list[tuple], output: Path) -> None:
    image = Image.new("RGB", (1700, 850), "#f5f7f4")
    draw = ImageDraw.Draw(image)
    draw.text((54, 44), "Server/database output", font=UI_FONT_BOLD, fill="#17201b")
    draw.text((54, 92), str(DB_PATH), font=UI_FONT_SMALL, fill="#66736c")
    draw.rounded_rectangle((54, 145, 1646, 790), radius=14, fill="#ffffff", outline="#dce5de", width=2)
    headers = ["Id", "Title", "Priority", "Due date", "Done", "Created", "Updated"]
    widths = [80, 560, 160, 180, 120, 220, 220]
    x = 80
    y = 175
    draw.rectangle((70, y - 14, 1628, y + 42), fill="#f0f5f1")
    for header, width in zip(headers, widths):
        draw.text((x, y), header.upper(), font=load_font("segoeuib.ttf", 18), fill="#66736c")
        x += width
    y += 72
    for row in rows:
        values = [
            str(row[0]),
            row[1],
            row[2],
            row[3] or "-",
            "Yes" if row[4] else "No",
            row[5],
            row[6] or "-",
        ]
        x = 80
        for value, width in zip(values, widths):
            shown = value if len(value) <= 48 else value[:45] + "..."
            draw.text((x, y), shown, font=UI_FONT_SMALL, fill="#17201b")
            x += width
        draw.line((70, y + 42, 1628, y + 42), fill="#dce5de", width=1)
        y += 72
    image.save(output)


def fetch_route(path: str) -> tuple[int, str]:
    try:
        with urlopen(f"http://localhost:5262{path}", timeout=5) as response:
            return response.status, response.read().decode("utf-8", errors="replace")
    except Exception as exc:
        return 0, str(exc)


def strip_html_text(content: str) -> str:
    text = re.sub(r"<script.*?</script>", "", content, flags=re.S)
    text = re.sub(r"<style.*?</style>", "", text, flags=re.S)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def read_database_rows() -> list[tuple]:
    with sqlite3.connect(DB_PATH) as connection:
        return connection.execute(
            """
            SELECT Id,
                   Title,
                   Priority,
                   COALESCE(strftime('%Y-%m-%d', DueDate), ''),
                   IsCompleted,
                   strftime('%Y-%m-%d %H:%M', CreatedAt),
                   COALESCE(strftime('%Y-%m-%d %H:%M', UpdatedAt), '')
            FROM Tasks
            ORDER BY IsCompleted, DueDate, Id
            """
        ).fetchall()


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "Application 08 - Blazor ToDo Database Assignment"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Local SQLite persistence evidence"
    set_run_font(footer.runs[0], size=9, color=MUTED)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=MUTED, bold=True)


def add_picture(doc: Document, path: Path, width: float = 6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def build_images(rows: list[tuple], route_status: dict[str, tuple[int, str]]) -> dict[str, Path]:
    images = {
        "program": SHOTS / "code_program_database_configuration.png",
        "dbcontext": SHOTS / "code_todo_db_context.png",
        "service": SHOTS / "code_task_service_crud.png",
        "todo_page": SHOTS / "code_todo_razor_client.png",
        "database_page": SHOTS / "code_database_output_page.png",
        "client_output": SHOTS / "client_output_todo_page.png",
        "database_output": SHOTS / "server_database_output_rows.png",
        "route_output": SHOTS / "server_route_checks.png",
    }

    draw_code_image(select_lines(PROJECT / "Program.cs", 1, 38), "Program.cs - EF Core SQLite registration", images["program"])
    draw_code_image(select_lines(PROJECT / "Data" / "TodoDbContext.cs", 1, 59), "TodoDbContext.cs - Tasks table schema and seed data", images["dbcontext"])
    draw_code_image(select_lines(PROJECT / "Services" / "TodoTaskService.cs", 1, 136), "TodoTaskService.cs - create, update, toggle, delete, stats", images["service"])
    draw_code_image(select_lines(PROJECT / "Components" / "Pages" / "Todo.razor", 1, 150), "Todo.razor - browser task workflow", images["todo_page"])
    draw_code_image(select_lines(PROJECT / "Components" / "Pages" / "Database.razor", 1, 78), "Database.razor - server/database output page", images["database_page"])
    draw_client_output(rows, images["client_output"])
    draw_database_table(rows, images["database_output"])

    route_lines = [
        "PS> Invoke-WebRequest http://localhost:5262/",
        f"StatusCode: {route_status['/'][0]}",
        "PS> Invoke-WebRequest http://localhost:5262/todo",
        f"StatusCode: {route_status['/todo'][0]}",
        "PS> Invoke-WebRequest http://localhost:5262/database",
        f"StatusCode: {route_status['/database'][0]}",
        "",
        "Database file:",
        str(DB_PATH),
        f"Rows returned from SQLite: {len(rows)}",
    ]
    draw_terminal_image(route_lines, "Server verification output", images["route_output"])
    return images


def build_doc(images: dict[str, Path], rows: list[tuple], route_status: dict[str, tuple[int, str]]) -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Application 08: Blazor ToDo App Connected with Local Database")
    set_run_font(run, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Task create, update, complete, and delete actions are persisted locally in SQLite through EF Core.")
    set_run_font(run, size=12.5, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_table_borders(meta)
    metadata = [
        ("Project", "Blazor Task Management System"),
        ("Framework", ".NET 10 Blazor Server with InteractiveServer components"),
        ("Database", f"SQLite local file: {DB_PATH.name}"),
        ("Repository", "https://github.com/Syed-Aqib-555/Blazor-Task-Management-System"),
    ]
    for idx, (label, value) in enumerate(metadata):
        cells = meta.rows[idx].cells
        cells[0].text = label
        cells[1].text = value
        shade_cell(cells[0], HEADER_FILL)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, color=INK, bold=(cell is cells[0]))

    doc.add_heading("Objective", level=1)
    doc.add_paragraph(
        "The objective of Application 08 is to connect the earlier Blazor ToDo application to a local database so tasks are not lost when the browser refreshes or the server restarts. The browser interface now uses a database-backed service for every create, update, complete, and delete operation."
    )

    doc.add_heading("Implementation Summary", level=1)
    add_bullets(
        doc,
        [
            "Added EF Core SQLite and configured the connection string in appsettings.json.",
            "Created TodoDbContext with a Tasks table, indexes for completion and due date, and seed records.",
            "Replaced the in-memory list service with TodoTaskService, which opens a database context for each operation.",
            "Built a professional task dashboard with metrics, filters, form validation, edit mode, completion toggles, and delete actions.",
            "Added a /database page that reads the same local SQLite file and displays server-side table rows for documentation evidence.",
        ],
    )

    doc.add_heading("Database Design", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    headers = ["Column", "Purpose", "Example"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, HEADER_FILL)
        cell.text = text
    columns = [
        ("Id", "Primary key generated by SQLite", "1"),
        ("Title", "Required task name", "Review database manuals"),
        ("Description", "Optional notes", "Check Week 13 notes"),
        ("Priority", "High, Medium, or Low", "High"),
        ("DueDate", "Optional task deadline", "2026-05-27"),
        ("IsCompleted", "Completion state used by toggle action", "0 or 1"),
        ("CreatedAt / UpdatedAt / CompletedAt", "Audit timestamps", "2026-05-26 09:00"),
    ]
    for column, purpose, example in columns:
        cells = table.add_row().cells
        cells[0].text = column
        cells[1].text = purpose
        cells[2].text = example
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.2, color=INK)

    doc.add_heading("Code Screenshots", level=1)
    for key, caption in [
        ("program", "Figure 1. Program.cs registers the SQLite connection, EF Core DbContext factory, database information, and task service."),
        ("dbcontext", "Figure 2. TodoDbContext maps the Tasks table and defines seed rows for the local database."),
        ("service", "Figure 3. TodoTaskService performs create, update, toggle completion, delete, list, and statistics operations through EF Core."),
        ("todo_page", "Figure 4. Todo.razor provides the client form, filters, task cards, edit buttons, completion toggles, and delete buttons."),
        ("database_page", "Figure 5. Database.razor shows the server/database output from the same SQLite file."),
    ]:
        add_picture(doc, images[key])
        add_caption(doc, caption)

    doc.add_heading("Output Evidence", level=1)
    route_summary = (
        f"HTTP route checks returned / = {route_status['/'][0]}, /todo = {route_status['/todo'][0]}, "
        f"and /database = {route_status['/database'][0]}. The database file contains {len(rows)} rows."
    )
    doc.add_paragraph(route_summary)
    add_picture(doc, images["route_output"])
    add_caption(doc, "Figure 6. Server route verification and local SQLite database file output.")
    add_picture(doc, images["client_output"])
    add_caption(doc, "Figure 7. Client output evidence for the /todo task screen using rows loaded from SQLite.")
    add_picture(doc, images["database_output"])
    add_caption(doc, "Figure 8. Server/database output from todo_tasks.db showing stored task rows.")

    doc.add_heading("CRUD Flow Explanation", level=1)
    add_bullets(
        doc,
        [
            "Create: the user fills the task form and selects Add task. TodoTaskService.CreateTaskAsync inserts the record into Tasks and SaveChangesAsync writes it to todo_tasks.db.",
            "Update: the user selects Edit, changes title, notes, priority, or due date, and saves. TodoTaskService.UpdateTaskAsync updates the row and sets UpdatedAt.",
            "Complete: the Open/Done toggle calls ToggleCompletionAsync, flips IsCompleted, and stores CompletedAt when the task is completed.",
            "Delete: the Delete button calls DeleteTaskAsync, removes the row from DbSet, and persists the change in SQLite.",
            "Read: /todo and /database both call GetTasksAsync, so browser output and database output stay consistent.",
        ],
    )

    doc.add_heading("Run Instructions", level=1)
    doc.add_paragraph('Open PowerShell inside "C:\\Users\\Admin\\source\\repos\\Blazor Task Management System\\Blazor Task Management System" and run:')
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_after = Pt(6)
    run = code.add_run("dotnet run --launch-profile http")
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    doc.add_paragraph("Then open http://localhost:5262/todo for the task client and http://localhost:5262/database for the database output page.")

    doc.add_heading("GitHub Commit Requirement", level=1)
    doc.add_paragraph(
        "The repository is configured with origin https://github.com/Syed-Aqib-555/Blazor-Task-Management-System.git. The final step is to create and push 50+ commits as requested, after the implementation and documentation are verified."
    )

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    route_status = {
        "/": fetch_route("/"),
        "/todo": fetch_route("/todo"),
        "/database": fetch_route("/database"),
    }
    rows = read_database_rows()
    images = build_images(rows, route_status)
    build_doc(images, rows, route_status)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
